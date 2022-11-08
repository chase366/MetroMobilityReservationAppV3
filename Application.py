# Sweigart, Al. Automate the Boring Stuff with Python: Practical Programming for Total Beginners. No Starch Press, 2015. 

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import sys
import os
from datetime import datetime
from datetime import date
from datetime import timedelta
import docx
from docx.shared import Pt

class WordDocument:

    def getDate(self, filename):
        doc = docx.Document(filename)
        dateParagraph = doc.paragraphs[8].text  # the paragraph where the date is
        indexOfColon = dateParagraph.find(":")  # gets the index where I begin to write in the date
        date_substring = dateParagraph[indexOfColon + 2:]  # the old date substring

        return date_substring

    def changeDate(self, newDate, filename):
        doc = docx.Document(filename)
        dateParagraph = doc.paragraphs[8].text  # the paragraph where the date is

        date_substring = self.getDate(filename) # the old date substring

        newString = dateParagraph.replace(date_substring, newDate) # stores a string that replaces the old date with the new date
        doc.paragraphs[8].text = newString # In the doc, replace the old string with the new string that has the new date
        doc.save(filename)

    def addDays(self, filename, numberOfDays):
        date_str = self.getDate(filename)
        datetime_obj = datetime.strptime(date_str, "%m/%d/%y")
        newDateTime = datetime_obj + timedelta(days=numberOfDays)
        newDate = newDateTime.date()
        newDateFormatted = newDate.strftime("%m/%d/%y")
        return str(newDateFormatted)

    def changeStyle(self, filename, paragraphNum):
        doc = docx.Document(filename)
        run = doc.paragraphs[8].add_run()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial Black"
        font.size = Pt(10)
        doc.save(filename)

class Email:

    my_email_address = ""
    recipient_email_address = ""
    password = ""
    subject = ""
    email_body = ""
    msg = MIMEMultipart()

    def __init__(self, my_email_address, password, recipient_email_address, subject, email_body):
        self.my_email_address = my_email_address
        self.recipient_email_address = recipient_email_address
        self.password = password
        self.subject = subject
        self.email_body = email_body

    def insert_attachment(self, filename):
        self.msg["From"] = self.my_email_address
        self.msg["To"] = self.recipient_email_address
        self.msg["Subject"] = self.subject

        self.msg.attach(MIMEText(self.email_body, "plain"))
        attachment = open(filename, "rb")

        p = MIMEBase("application", "octet-stream")
        p.set_payload(attachment.read())

        encoders.encode_base64(p)
        p.add_header("Content-Disposition", "attachment; filename= %s" % filename)

        self.msg.attach(p)

    def send(self):
        smtpObj = smtplib.SMTP("smtp.comcast.net", 587) # Email protocol
        smtpObj.starttls() # Enable encryption
        smtpObj.login(self.my_email_address, self.password)
        sent_mail = False

        text = self.msg.as_string()
        dict = smtpObj.sendmail(self.my_email_address, self.recipient_email_address, text)

        # If dict is empty, then sending the e-mail was successful
        if len(dict) == 0:
            sent_mail = True

        smtpObj.quit() # Disconnect from e-mail server

        return sent_mail

    # Getters and setters
    def setMyEmailAddress(self, my_email_address):
        self.my_email_address = my_email_address

    def setRecipientEmailAddress(self, recipient_email_address):
        self.recipient_email_address = recipient_email_address

    def setPassword(self, password):
        self.password = password

    def setSubject(self, subject):
        self.subject = subject

    def setEmailBody(self, email_body):
        self.email_body = email_body

    def getMyEmailAddress(self):
        return self.my_email_address

    def getRecipientEmailAddress(self):
        return self.recipient_email_address

    def getSubject(self):
        return self.subject

    def getEmailBody(self):
        return self.email_body

    def getPassword(self):
        return self.password

# Add 7 days (a week) from the current dates to each of the forms

filename = "email_reservation_north_hennepin_Mondays.docx"
doc = WordDocument()
newDateStr = doc.addDays(filename, 7)
doc.changeDate(newDateStr, filename)

filename2 = "email_reservation_north_hennepin_Wednesdays.docx"
doc = WordDocument()
newDateStr = doc.addDays(filename2, 7)
doc.changeDate(newDateStr, filename2)


filename = "email_reservation_grandma.docx"
doc = WordDocument()
newDateStr = doc.addDays(filename3, 7)
doc.changeDate(newDateStr, filename3)

my_email_address = "chaseconner@comcast.net"
password = "h0X$0%2JN5Z24mL4"
recipient_email_address = "WReservations@metc.state.mn.us"

subject = "Email Reservation Forms Attached"
email_body = "Dear Metro Mobility,\n\nI have attached my e-mail reservation form(s)" \
             " to this e-mail. You may call me at 763-202-1523 to let me" \
             " know the pickup time(s) the day before. You may also give my " \
             "number to the bus driver in case he/she has trouble locating " \
             "me.\n\nThank you,\nChase Conner"

email = Email(my_email_address, password, recipient_email_address, subject, email_body)

# email.insert_attachment(filename)
# email.insert_attachment(filename2)
email.insert_attachment(filename3)

print("Attempting to send e-mail...")
sent_mail = email.send()

now = datetime.now()
dt_string = now.strftime("%m/%d/%Y %H:%M:%S")

f = open("log.txt", "a")

if sent_mail == True:
    # print("\nE-mail to " + email.getRecipientEmailAddress() + " was successfully sent on " + dt_string)
    f.write("E-mail to " + email.getRecipientEmailAddress() + " was successfully sent on " + dt_string + "\n")
else:
    # print("\nE-mail to " + email.getRecipientEmailAddress() + " failed to send on " + dt_string)
    f.write("E-mail to " + email.getRecipientEmailAddress() + " failed to send on " + dt_string + "\n")

f.close()

#input("\nPress [ENTER] to exit")
#sys.exit()